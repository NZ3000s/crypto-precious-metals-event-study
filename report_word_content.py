#!/usr/bin/env python3
"""
Generate report.docx: report that follows the presentation structure, with all
figures (and captions) and tables. Run after: python3 export_presentation_data.py
Requires: pip install python-docx
"""
import os
import json

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Install python-docx: pip install python-docx")
    raise

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(PROJECT_ROOT, "report.docx")
FIG_DIR = os.path.join(PROJECT_ROOT, "presentation_figures")
EMBED_JS = os.path.join(PROJECT_ROOT, "presentation_data_embed.js")

# Figures in order of appearance (filename, caption)
FIGURES = [
    ("why_story_two_venues.png", "Figure 1: Before the event — one venue (ETFs only); after the event — two venues (ETFs + Binance perpetuals)."),
    ("event_method_design.png", "Figure 2: Event-study design — timeline (pre / event / post) and OLS specification (outcome = α + β·post_event + γ·controls)."),
    ("cumulative_returns.png", "Figure 3: Cumulative returns (index 100). ETFs from first date; Binance perpetuals from launch date only. Silver (SLV, XAGUSDT) outperforms gold in this sample; all four series move together after the event."),
    ("vol_volume.png", "Figure 4: Realized 20d volatility, Parkinson volatility, and ETF volume over time. Vertical line = event (2026-01-05). Silver (SLV) reacts more strongly post-event."),
    ("pre_post_bars.png", "Figure 5: Percent change (post vs pre) for volatility and volume. Silver increases more than gold; same units (%) make the comparison clear."),
    ("correlation_matrix.png", "Figure 6: Correlation matrix of key variables — returns (GLD, SLV, XAU, XAG), macro controls (DXY, VIX, 10Y), and realized volatility."),
    ("rolling_corr.png", "Figure 7: Rolling 20-day correlation between ETF and Binance perpetual returns (GLD–XAU, SLV–XAG). Data only after launch; first point ≈20 trading days (rolling window). High values = integration."),
]


def load_key_table():
    """Load keyTable from presentation_data_embed.js if it exists."""
    if not os.path.exists(EMBED_JS):
        return None
    try:
        with open(EMBED_JS, encoding="utf-8") as f:
            raw = f.read()
        # JSON is the first { ... }; extract by brace matching
        start = raw.find("{")
        if start == -1:
            return None
        depth = 0
        for i in range(start, len(raw)):
            if raw[i] == "{":
                depth += 1
            elif raw[i] == "}":
                depth -= 1
                if depth == 0:
                    data = json.loads(raw[start : i + 1])
                    return data.get("keyTable")
    except Exception:
        pass
    return None


def add_figure(doc, rel_path, caption, width_inches=5.5):
    """Add an image and caption; if file missing, add placeholder."""
    path = os.path.join(PROJECT_ROOT, rel_path)
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(width_inches))
            p = doc.add_paragraph()
            p.add_run(caption).italic = True
            p.paragraph_format.space_before = Pt(6)
        except Exception:
            doc.add_paragraph(f"[Image not loaded: {rel_path}]")
            doc.add_paragraph(caption)
    else:
        doc.add_paragraph(f"[Run python3 export_presentation_data.py to generate {rel_path}]")
        doc.add_paragraph(caption)
    doc.add_paragraph()


def add_table_from_rows(doc, headers, rows, style="Table Grid"):
    """Add a Word table with headers and rows. Each row is a dict keyed by header, or a list."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = style
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = str(h)
    for i, row in enumerate(rows):
        for j, h in enumerate(headers):
            if isinstance(row, dict):
                val = row.get(h, "—")
            else:
                val = row[j] if j < len(row) else "—"
            table.rows[i + 1].cells[j].text = str(val) if val is not None else "—"
    doc.add_paragraph()


def main():
    doc = Document()
    key_table = load_key_table()

    # ----- Title & Abstract -----
    title = doc.add_heading("The Impact of Binance Gold and Silver Perpetuals on ETF Markets: An Event Study", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Abstract. ").bold = True
    p.add_run(
        "We study whether the launch of Binance TRADFI gold and silver perpetual futures (XAUUSDT, XAGUSDT) on 5 January 2026 "
        "changed the volatility and liquidity of established gold and silver ETFs (GLD, SLV). Using a daily panel and an event-study design "
        "with pre- and post-event comparison, we find that ETF volatility and trading volume increased after the launch. "
        "When we control for macro factors (dollar index, VIX, and 10-year yield), the post-event effect remains positive and significant. "
        "Rolling correlation between ETF and Binance perpetual returns is high after the launch, consistent with price integration. "
        "Limitations include the absence of pre-event Binance data and the short post-event sample."
    )
    doc.add_paragraph()

    # ----- 1. Research question -----
    doc.add_heading("1. Research question", level=1)
    doc.add_paragraph(
        "Did the introduction of Binance TRADFI gold and silver perpetual futures (XAUUSDT, XAGUSDT) change: (1) the volatility of gold and silver ETFs (GLD, SLV); "
        "(2) liquidity (volume, Amihud illiquidity); (3) price integration between ETF and perpetual? We have integration data only after the launch (rolling correlation GLD–XAU, SLV–XAG); "
        "we do not have pre-event Binance data, so we describe post-event integration only."
    )
    doc.add_paragraph("Event date: 2026-01-05. We test two hypotheses: H1 — ETF volatility increases after the launch; H2 — volume/liquidity changes (we expect volume to rise).")
    doc.add_paragraph()

    # ----- 2. Why this story matters -----
    doc.add_heading("2. Why this story matters", level=1)
    doc.add_paragraph(
        "Gold and silver are traded as ETFs (GLD, SLV) on regulated exchanges. In early 2026, Binance started listing perpetuals on the same assets (XAUUSDT, XAGUSDT). "
        "So there are two venues for exposure to the same thing. We ask: did the ETF market change after that? More volatile? More trading? More linked to the new venue? "
        "This matters for traders, market-makers, and regulators. Augustin et al. (2023) show that the introduction of derivatives can change spot markets—volatility, liquidity, and how prices move together."
    )
    add_figure(doc, os.path.join("presentation_figures", FIGURES[0][0]), FIGURES[0][1])
    doc.add_paragraph()

    # ----- 3. Data: sources and periods -----
    doc.add_heading("3. Data: sources and periods", level=1)
    headers_data = ["Source", "What we get", "Frequency", "Rows", "Note"]
    rows_data = [
        ["Binance API", "XAUUSDT, XAGUSDT (OHLCV), funding, OI", "1m", "67k+", "From 2026-01-05 (XAU), 2026-01-07 (XAG)"],
        ["Yahoo Finance", "ETFs: GLD, SLV. Controls: DXY, VIX, TNX. GLD/SLV 1m", "Daily; 1m last 7d", "118 daily", "1m only last ~7 days"],
        ["FRED", "DGS10 — 10Y Treasury (macro control)", "Daily", "in daily panel", "Same role as TNX"],
    ]
    add_table_from_rows(doc, headers_data, rows_data)
    doc.add_paragraph(
        "Panels: Daily = 118 rows (event study, OLS). 1-minute = 67k+ rows (intraday analysis). Integration: post-event only (no Binance before 2026-01-05). "
        "DXY = dollar index; VIX = volatility index; TNX / DGS10 = 10Y yield. Used as controls in regressions."
    )
    doc.add_paragraph()

    # ----- 4. How we use the data -----
    doc.add_heading("4. How we use the data", level=1)
    doc.add_paragraph(
        "Daily panel (features_daily.csv): 118 days × many variables. Used for event study—pre/post means, t-tests, OLS with post_event and controls (DXY, VIX, 10Y). "
        "1-minute panel (features_1m.csv): one row per Binance XAUUSDT 1m candle, daily ETF/controls merged by date; used for intraday return distribution and daily realized vol from 1m. "
        "Minute-level correlation: last 7 days where Yahoo 1m (GLD/SLV) overlaps with Binance 1m—same-frequency correlation GLD–XAU, SLV–XAG. "
        "Integration (descriptive): rolling correlation in the daily panel; we report post-event level only."
    )
    doc.add_paragraph()

    # ----- 5. Event and method -----
    doc.add_heading("5. Event and method", level=1)
    doc.add_paragraph(
        "Design: Event study with pre/post comparison. No control group → we do not use difference-in-differences. "
        "Tests: (1) Welch t-tests (pre vs post); (2) OLS with outcome = volatility or liquidity, regressor = post_event + controls (DXY, VIX, 10Y). HAC (Newey–West) SE; winsorization 5%/95%. "
        "Why controls? DXY, VIX, 10Y capture general market conditions. By including them, the coefficient on post_event measures the effect of the listing itself, not just the market. "
        "Literature: Augustin et al. (2023) use DiD with treatment/control; we have a single market before/after, so we use pre/post + OLS with controls."
    )
    add_figure(doc, os.path.join("presentation_figures", FIGURES[1][0]), FIGURES[1][1])
    doc.add_paragraph()

    # ----- 6. Cumulative returns -----
    doc.add_heading("6. Cumulative returns (ETFs vs Binance)", level=1)
    doc.add_paragraph(
        "On the chart, higher line = higher cumulative return. Silver (SLV, XAGUSDT) outperforms gold (GLD, XAUUSDT) in this sample. "
        "After the event, all four series move together—that co-movement is the first visual evidence of integration. Binance series are plotted only from the launch date (no data before)."
    )
    add_figure(doc, os.path.join("presentation_figures", FIGURES[2][0]), FIGURES[2][1], width_inches=6)
    doc.add_paragraph()

    # ----- 7. Volatility and volume over time -----
    doc.add_heading("7. Volatility and volume over time", level=1)
    doc.add_paragraph(
        "Volatility and volume dynamics around the event—basis for H1 and H2. Silver (SLV) reacts more strongly: post-event, both realized and Parkinson volatility, and ETF volume, rise more for SLV than for GLD. "
        "Chart titles show percent growth (post vs pre) for each series."
    )
    add_figure(doc, os.path.join("presentation_figures", FIGURES[3][0]), FIGURES[3][1], width_inches=6)
    doc.add_paragraph()

    # ----- 8. Pre vs post: how much did each metric grow? -----
    doc.add_heading("8. Pre vs post: how much did each metric grow?", level=1)
    doc.add_paragraph(
        "After the event, volatility and volume increase (positive %); silver (SLV) increases more than gold (GLD). We test whether these changes are statistically significant (two-sample Welch t-test: H₀ mean(post)=mean(pre)). "
        "Low p-value (e.g. p < 0.05) → reject H₀ → change is significant. We then control for macro in regressions."
    )
    add_figure(doc, os.path.join("presentation_figures", FIGURES[4][0]), FIGURES[4][1], width_inches=5.5)
    if key_table and len(key_table) >= 4:
        doc.add_paragraph("Table 1: Pre/post means, percent change, t-statistic, and t-test p-value (first four metrics).")
        t_headers = ["Metric", "Pre mean", "Post mean", "Pct", "t", "t-test p"]
        add_table_from_rows(doc, t_headers, [dict((k, r.get(k, "—")) for k in t_headers) for r in key_table[:4]])
    else:
        doc.add_paragraph("Table 1 (t-test): Run export_presentation_data.py to generate presentation_data_embed.js; then re-run this script.")
    doc.add_paragraph()

    # ----- 9. Correlation matrix -----
    doc.add_heading("9. Correlation matrix (key variables)", level=1)
    doc.add_paragraph(
        "Pairwise correlations between returns (GLD, SLV, XAU, XAG), macro controls (DXY, VIX, 10Y/TNX), and realized volatility. "
        "Red = positive, blue = negative. High correlation between GLD and XAU (and SLV and XAG) in the post-event overlap supports integration; macro variables are used as controls in regressions."
    )
    add_figure(doc, os.path.join("presentation_figures", FIGURES[5][0]), FIGURES[5][1], width_inches=5.5)
    doc.add_paragraph()

    # ----- 10. Integration: rolling correlation -----
    doc.add_heading("10. Integration: do ETF and Binance perpetual move together?", level=1)
    doc.add_paragraph(
        "Rolling 20-day correlation between ETF returns and Binance perpetual returns (GLD–XAU, SLV–XAG). We have data only after the launch; the first point appears ≈20 trading days after the event (rolling window). "
        "Values near 0.9–1.0 mean the two venues move in lockstep—one price-discovery process. We cannot test 'integration increased' because we have no pre-event Binance data; we describe the post-event level."
    )
    add_figure(doc, os.path.join("presentation_figures", FIGURES[6][0]), FIGURES[6][1], width_inches=6)
    doc.add_paragraph()

    # ----- 11. Regression results -----
    doc.add_heading("11. Regression results", level=1)
    doc.add_paragraph(
        "Key result: Positive and significant coefficient on post_event in volatility regressions (GLD and SLV) when we control for DXY, VIX, 10Y → H1 supported. "
        "So the rise in volatility is attributed to the listing, not to general market moves. Volume higher post-event → H2 supported (direction). "
        "The table below reports pre- and post-event means, percent change, t-test (t-statistic and p-value), the OLS coefficient on post_event, and the regression p-value for that coefficient (HAC standard errors). "
        "Corr GLD–XAU and Corr SLV–XAG have no pre-event values (—) because Binance did not list before the event; their post-event means are our price-integration summary. "
        "Higher ETF volatility need not mean chaos—it can reflect more hedging and informed trading across the two venues."
    )
    if key_table:
        doc.add_paragraph("Table 2: Key results — pre/post means, percent change, t-test, OLS post_event coefficient and regression p-value.")
        full_headers = ["Metric", "Pre mean", "Post mean", "Pct", "t", "t-test p", "post coef", "reg p"]
        add_table_from_rows(doc, full_headers, [dict((k, r.get(k, "—")) for k in full_headers) for r in key_table])
    else:
        doc.add_paragraph("Table 2 (key results): Run export_presentation_data.py, then re-run this script.")
    doc.add_paragraph()

    # ----- 12. Conclusions -----
    doc.add_heading("12. Conclusions", level=1)
    doc.add_paragraph(
        "The listing of Binance gold/silver perpetuals is associated with higher ETF volatility and volume (H1, H2 supported when we control for macro). "
        "The two venues move together after the launch—integration is high where we can measure it. New derivatives on a new venue can shift incumbent markets; the story is similar to Bitcoin futures (Augustin et al.)."
    )
    doc.add_paragraph(
        "Findings: Volatility and volume up post-event (t-tests and OLS with DXY, VIX, 10Y). Rolling correlation GLD–XAU, SLV–XAG high after launch. "
        "Who cares: ETF market-makers, arbitrageurs, regulators—two venues, one price-discovery process."
    )
    doc.add_paragraph(
        "Limitations: No pre-event Binance data (no 'integration increased' test); no control group (no DiD); short post window—more data would help. "
        "Despite these caveats, the consistency of t-tests and OLS with controls supports the view that the listing is associated with the observed changes in ETF markets."
    )
    doc.add_paragraph()

    # ----- References -----
    doc.add_heading("References", level=1)
    doc.add_paragraph("Augustin, P., Rubtsov, A., & Shin, D. (2023). The impact of Bitcoin futures on spot markets. Management Science.")
    doc.add_paragraph()

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Approximate word count (body): {word_count}")
    print("Ensure presentation_figures/*.png exist (run python3 export_presentation_data.py first).")


if __name__ == "__main__":
    main()
